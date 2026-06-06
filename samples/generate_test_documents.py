from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "test-documents"
FONT_DIR = Path(r"C:\Windows\Fonts")
BODY_FONT = FONT_DIR / "simsun.ttc"
HEAD_FONT = FONT_DIR / "simhei.ttf"


def slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    return text.strip("-")


def set_run_font(run, font_name: str = "SimSun", size: int | None = None, bold: bool | None = None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH | None = None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, "SimSun", 9, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_doc_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def add_docx_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    if style and style.startswith("Heading"):
        set_run_font(run, "SimHei")
    else:
        set_run_font(run, "SimSun")
    return paragraph


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_text(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_margins(cells[idx])
    doc.add_paragraph()
    return table


def add_list(doc: Document, items: list[str], ordered: bool = False):
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        run = paragraph.add_run(item)
        set_run_font(run, "SimSun")


def build_docx(sample: dict):
    doc = Document()
    setup_doc_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run(sample["title"])
    set_run_font(title_run, "SimHei", 20, True)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(sample["subtitle"])
    set_run_font(subtitle_run, "SimSun", 10)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    for block in sample["blocks"]:
        kind = block["type"]
        if kind == "h1":
            add_docx_paragraph(doc, block["text"], "Heading 1")
        elif kind == "h2":
            add_docx_paragraph(doc, block["text"], "Heading 2")
        elif kind == "h3":
            add_docx_paragraph(doc, block["text"], "Heading 3")
        elif kind == "p":
            add_docx_paragraph(doc, block["text"])
        elif kind == "bullets":
            add_list(doc, block["items"], ordered=False)
        elif kind == "steps":
            add_list(doc, block["items"], ordered=True)
        elif kind == "table":
            add_docx_table(doc, block["headers"], block["rows"], block["widths"])
        elif kind == "page_break":
            doc.add_page_break()

    path = OUT_DIR / f"{sample['file_stem']}.docx"
    doc.save(path)
    return path


def register_pdf_fonts():
    body_path = BODY_FONT if BODY_FONT.exists() else FONT_DIR / "simsunb.ttf"
    head_path = HEAD_FONT if HEAD_FONT.exists() else FONT_DIR / "simhei.ttf"
    pdfmetrics.registerFont(TTFont("CJKBody", str(body_path)))
    pdfmetrics.registerFont(TTFont("CJKHead", str(head_path)))


def pdf_styles():
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CJKBody",
        parent=styles["BodyText"],
        fontName="CJKBody",
        fontSize=10.5,
        leading=15,
        alignment=TA_LEFT,
        wordWrap="CJK",
        spaceAfter=7,
    )
    h1 = ParagraphStyle(
        "CJKH1",
        parent=body,
        fontName="CJKHead",
        fontSize=17,
        leading=22,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "CJKH2",
        parent=body,
        fontName="CJKHead",
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=9,
        spaceAfter=6,
    )
    h3 = ParagraphStyle(
        "CJKH3",
        parent=body,
        fontName="CJKHead",
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=8,
        spaceAfter=4,
    )
    title = ParagraphStyle(
        "CJKTitle",
        parent=body,
        fontName="CJKHead",
        fontSize=20,
        leading=26,
        spaceAfter=4,
    )
    meta = ParagraphStyle(
        "CJKMeta",
        parent=body,
        fontName="CJKBody",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#555555"),
        spaceAfter=12,
    )
    return {"title": title, "meta": meta, "body": body, "h1": h1, "h2": h2, "h3": h3}


def esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(sample: dict):
    register_pdf_fonts()
    styles = pdf_styles()
    path = OUT_DIR / f"{sample['file_stem']}.pdf"
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    story = [
        Paragraph(esc(sample["title"]), styles["title"]),
        Paragraph(esc(sample["subtitle"]), styles["meta"]),
    ]

    for block in sample["blocks"]:
        kind = block["type"]
        if kind in ("h1", "h2", "h3"):
            story.append(Paragraph(esc(block["text"]), styles[kind]))
        elif kind == "p":
            story.append(Paragraph(esc(block["text"]), styles["body"]))
        elif kind == "bullets":
            for item in block["items"]:
                story.append(Paragraph(f"- {esc(item)}", styles["body"]))
        elif kind == "steps":
            for idx, item in enumerate(block["items"], start=1):
                story.append(Paragraph(f"{idx}. {esc(item)}", styles["body"]))
        elif kind == "table":
            data = [[Paragraph(esc(cell), styles["body"]) for cell in block["headers"]]]
            data.extend([[Paragraph(esc(cell), styles["body"]) for cell in row] for row in block["rows"]])
            width_total = 6.5 * inch
            col_widths = [width_total * width / sum(block["widths"]) for width in block["widths"]]
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), "CJKBody"),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D7DE")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 8))
        elif kind == "page_break":
            story.append(PageBreak())

    doc.build(story)
    return path


SAMPLES = [
    {
        "file_stem": "01-wenqu-architecture-deployment-guide",
        "title": "问渠企业智能知识库问答系统：架构与部署说明",
        "subtitle": "测试用途：标题层级、组件表格、端口信息、租户隔离规则。版本：v1.0",
        "blocks": [
            {"type": "h1", "text": "1. 系统定位"},
            {
                "type": "p",
                "text": "问渠是面向企业内部知识库的问答系统，支持文档上传、解析、向量化、检索增强生成和评估闭环。系统以多租户为基础隔离单位，所有业务数据、向量集合和对象存储桶都必须按租户划分。",
            },
            {"type": "h2", "text": "1.1 技术栈白名单"},
            {
                "type": "table",
                "headers": ["层次", "技术", "固定约束"],
                "rows": [
                    ["前端", "Vue3 + Vite + Element Plus", "Vue 3.4+，不得替换为 React"],
                    ["网关", "Spring Cloud Gateway 4.x", "负责 JWT 校验、限流和请求头注入"],
                    ["主服务", "Spring Boot 3.2+ / Java 17", "单体服务，包含用户、文档、对话、管理和评估模块"],
                    ["AI 框架", "Spring AI 1.0+", "连接 DeepSeek Chat 与智谱 embedding-3"],
                    ["向量库", "Milvus 2.4+", "每个租户一个 Collection，命名 tenant_{tenantId}_docs"],
                ],
                "widths": [1.1, 2.5, 2.9],
            },
            {"type": "h2", "text": "1.2 Docker Compose 服务"},
            {
                "type": "table",
                "headers": ["服务名", "内部端口", "职责"],
                "rows": [
                    ["nginx", "80", "唯一对外入口，转发到 frontend 和 gateway"],
                    ["gateway", "8080", "注入 X-User-Id 与 X-Tenant-Id"],
                    ["app", "8081", "Java 主服务，处理核心业务接口"],
                    ["parser", "8090", "Python FastAPI 文档解析侧车"],
                    ["ragas-svc", "8091", "Python FastAPI Ragas 评估侧车"],
                    ["milvus", "19530", "存储 2048 维 embedding 向量"],
                    ["minio", "9000", "按 tenant-{tenantId} Bucket 保存原始文件"],
                ],
                "widths": [1.2, 1.0, 4.3],
            },
            {"type": "h1", "text": "2. 多租户隔离"},
            {
                "type": "bullets",
                "items": [
                    "数据库查询通过 MyBatis-Plus TenantLineInnerInterceptor 自动追加 tenant_id 条件。",
                    "向量数据按 tenant_{tenantId}_docs Collection 隔离，不允许跨租户检索。",
                    "对象存储 Bucket 使用 tenant-{tenantId}，名称为小写短横线形式。",
                    "业务代码从 TenantContext 的 ThreadLocal 读取 tenantId 与 userId。",
                ],
            },
            {"type": "h1", "text": "3. 推荐验收问题"},
            {
                "type": "bullets",
                "items": [
                    "问渠系统的向量库使用什么技术？",
                    "parser 服务监听哪个端口？",
                    "为什么每个租户要有独立的 Milvus Collection？",
                ],
            },
        ],
    },
    {
        "file_stem": "02-wenqu-user-admin-manual",
        "title": "问渠用户与管理员操作手册",
        "subtitle": "测试用途：步骤分块、权限说明、API 路径、任务状态轮询。版本：v1.0",
        "blocks": [
            {"type": "h1", "text": "1. 角色与权限"},
            {
                "type": "table",
                "headers": ["角色", "标识", "主要权限"],
                "rows": [
                    ["超级管理员", "SUPER_ADMIN", "创建或禁用租户，查看所有租户数据"],
                    ["租户管理员", "TENANT_ADMIN", "管理本租户用户，上传或删除文档，运行评估"],
                    ["普通用户", "USER", "在本租户内进行问答，查看自己的会话历史"],
                ],
                "widths": [1.4, 1.5, 3.6],
            },
            {"type": "h1", "text": "2. 登录与会话"},
            {
                "type": "steps",
                "items": [
                    "用户访问登录页，提交 tenantCode、username 和 password。",
                    "网关转发 POST /api/auth/login 请求，主服务返回 accessToken 与 refreshToken。",
                    "前端后续请求携带 Authorization: Bearer {jwt}。",
                    "网关校验 JWT 后注入 X-User-Id 与 X-Tenant-Id。",
                    "用户可以调用 POST /api/chat/conversations 创建新对话。",
                ],
            },
            {"type": "h2", "text": "2.1 SSE 问答接口"},
            {
                "type": "p",
                "text": "POST /api/chat/conversations/{id}/ask 使用 text/event-stream。服务端以 event:token 逐字推送模型回答，完成后以 event:done 推送 source_chunks JSON，用于前端展示引用来源。",
            },
            {"type": "h1", "text": "3. 文档上传流程"},
            {
                "type": "steps",
                "items": [
                    "租户管理员在文档管理页上传 pdf 或 docx 文件。",
                    "接口 POST /api/docs/upload 立即返回 docId 与 taskId。",
                    "前端轮询 GET /api/tasks/{taskId}/status 展示进度。",
                    "后台异步任务依次执行 MinIO 存储、Python 解析、Embedding、Milvus 写入和 doc_chunk 入库。",
                    "当进度为 100 且状态为 DONE 时，文档可以参与问答检索。",
                ],
            },
            {"type": "h2", "text": "3.1 状态码解释"},
            {
                "type": "table",
                "headers": ["状态", "说明", "前端建议"],
                "rows": [
                    ["PENDING", "任务已创建但尚未开始", "显示排队中"],
                    ["RUNNING", "后台正在处理", "显示进度条"],
                    ["DONE", "任务已完成", "允许进入详情或问答"],
                    ["FAILED", "任务失败，errorMsg 存在原因", "显示重试或联系管理员入口"],
                ],
                "widths": [1.2, 3.2, 2.1],
            },
            {"type": "h1", "text": "4. 推荐验收问题"},
            {
                "type": "bullets",
                "items": [
                    "普通用户是否可以上传文档？",
                    "文档上传接口返回哪些字段？",
                    "问答接口的 done 事件会携带什么数据？",
                ],
            },
        ],
    },
    {
        "file_stem": "03-enterprise-policy-knowledge-base-sample",
        "title": "企业制度知识库样例：差旅、报销与休假",
        "subtitle": "测试用途：业务制度问答、数字检索、表格抽取、来源引用展示。版本：2026-06",
        "blocks": [
            {"type": "h1", "text": "1. 差旅申请"},
            {
                "type": "p",
                "text": "员工因公出差前，应至少提前 2 个工作日在 OA 系统提交差旅申请。申请内容包括出差目的、预计行程、预算金额、同行人员和客户联系人。未经审批的差旅费用原则上不予报销。",
            },
            {"type": "h2", "text": "1.1 住宿标准"},
            {
                "type": "table",
                "headers": ["城市类别", "普通员工上限", "部门负责人上限", "备注"],
                "rows": [
                    ["一线城市", "600 元/晚", "800 元/晚", "北京、上海、广州、深圳"],
                    ["省会城市", "450 元/晚", "650 元/晚", "含直辖市以外省会"],
                    ["其他城市", "350 元/晚", "500 元/晚", "按实际票据报销"],
                ],
                "widths": [1.2, 1.4, 1.5, 2.4],
            },
            {"type": "h2", "text": "1.2 交通标准"},
            {
                "type": "bullets",
                "items": [
                    "高铁行程 5 小时以内应优先选择二等座。",
                    "机票应选择经济舱，除非客户日程要求导致无合理替代航班。",
                    "市内交通应优先使用公共交通或合规网约车。",
                ],
            },
            {"type": "h1", "text": "2. 报销材料"},
            {
                "type": "p",
                "text": "报销人应在差旅结束后 10 个自然日内提交报销单，并上传发票、行程单、酒店水单和审批截图。财务部在资料齐全后 5 个工作日内完成审核。",
            },
            {"type": "h2", "text": "2.1 不予报销事项"},
            {
                "type": "bullets",
                "items": [
                    "与出差目的无关的个人消费。",
                    "超出标准且未提前审批的住宿或交通费用。",
                    "票据抬头、税号或日期明显错误且无法补正的费用。",
                ],
            },
            {"type": "h1", "text": "3. 年假与病假"},
            {
                "type": "table",
                "headers": ["假期类型", "申请提前量", "证明材料", "审批人"],
                "rows": [
                    ["年假", "至少提前 3 个工作日", "无特殊材料", "直属经理"],
                    ["病假", "当天或返岗后 2 日内", "医院证明或病历", "直属经理与 HR"],
                    ["婚假", "至少提前 10 个工作日", "结婚证复印件", "直属经理与 HR"],
                ],
                "widths": [1.2, 1.8, 2.0, 1.5],
            },
            {"type": "h1", "text": "4. 推荐验收问题"},
            {
                "type": "bullets",
                "items": [
                    "一线城市普通员工住宿报销上限是多少？",
                    "差旅结束后多久内要提交报销单？",
                    "申请年假需要提前几个工作日？",
                ],
            },
        ],
    },
]


def build_scanned_pdf():
    pdf_path = OUT_DIR / "04-scanned-ocr-notice.pdf"
    stale_image_path = OUT_DIR / "04-scanned-ocr-notice-page.png"
    if stale_image_path.exists():
        stale_image_path.unlink()
    width, height = 1240, 1754
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(str(HEAD_FONT), 54)
    body_font = ImageFont.truetype(str(BODY_FONT), 34)
    small_font = ImageFont.truetype(str(BODY_FONT), 28)

    draw.rectangle((70, 70, width - 70, height - 70), outline=(180, 180, 180), width=3)
    draw.text((160, 150), "问渠知识库扫描件 OCR 测试通知", font=title_font, fill=(20, 20, 20))
    lines = [
        "编号：WQ-OCR-2026-001",
        "",
        "请各部门在上传纸质制度扫描件前确认页面清晰、方向正确。",
        "若 PDF 页面可提取字符数低于 100，解析侧车会进入 OCR 流程。",
        "扫描件建议分辨率不低于 300 DPI，避免阴影、折痕和遮挡。",
        "",
        "测试问题：扫描件的建议分辨率是多少？",
        "测试问题：低于多少字符会判定为扫描版？",
    ]
    y = 280
    for line in lines:
        draw.text((150, y), line, font=body_font if line else small_font, fill=(40, 40, 40))
        y += 58 if line else 38

    draw.text((150, height - 210), "信息安全部 / 2026 年 6 月", font=small_font, fill=(90, 90, 90))
    image.save(pdf_path, "PDF", resolution=300.0)
    return pdf_path


def write_readme(files: list[str]):
    readme = OUT_DIR / "README.md"
    content = """# 问渠测试文档样本

这些文件用于测试 `kb-parser` 的 PDF/Word 解析、标题分块、段落分块、表格抽取、OCR 扫描件识别，以及 RAG 检索引用展示。

## 文件清单

"""
    for file_name in files:
        content += f"- `{file_name}`\n"
    content += """
## 建议测试问题

- 问渠系统的向量库使用什么技术？
- parser 服务监听哪个端口？
- 文档上传接口返回哪些字段？
- 问答接口的 done 事件会携带什么数据？
- 一线城市普通员工住宿报销上限是多少？
- 差旅结束后多久内要提交报销单？
- 扫描件的建议分辨率是多少？
"""
    readme.write_text(content, encoding="utf-8")
    return readme


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = []
    manifest = []

    for sample in SAMPLES:
        docx_path = build_docx(sample)
        pdf_path = build_pdf(sample)
        generated.extend([docx_path.name, pdf_path.name])
        manifest.append(
            {
                "title": sample["title"],
                "docx": docx_path.name,
                "pdf": pdf_path.name,
                "suggested_questions": [
                    item
                    for block in sample["blocks"]
                    if block["type"] == "bullets"
                    for item in block["items"]
                    if item.endswith("？")
                ],
            }
        )

    scanned_pdf_path = build_scanned_pdf()
    generated.append(scanned_pdf_path.name)
    manifest.append(
        {
            "title": "问渠知识库扫描件 OCR 测试通知",
            "pdf": scanned_pdf_path.name,
            "note": "图片型 PDF，用于触发扫描件 OCR 路径。",
            "suggested_questions": ["扫描件的建议分辨率是多少？", "低于多少字符会判定为扫描版？"],
        }
    )

    write_readme(generated)
    (OUT_DIR / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"Generated {len(generated)} files in {OUT_DIR}")
    for name in generated:
        print(name)


if __name__ == "__main__":
    main()
